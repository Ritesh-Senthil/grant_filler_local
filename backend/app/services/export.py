from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from app.models import Answer, Grant, Question
from app.services.export_datetime import format_export_timestamp


@dataclass(frozen=True)
class ExportContext:
    """Per-export metadata (locale-formatted time, optional org branding)."""

    exported_at_label: str
    organization_line: str | None = None


def _default_context() -> ExportContext:
    now = datetime.now(timezone.utc)
    return ExportContext(
        exported_at_label=format_export_timestamp(now, "iso"),
        organization_line=None,
    )


def _txt(s: str) -> str:
    """fpdf2 core fonts are latin-1; normalize for safety."""
    if not s:
        return ""
    return s.encode("latin-1", errors="replace").decode("latin-1")


def _format_answer_value(a: Answer | None) -> str:
    if a is None or a.answer_value is None:
        return "(no answer yet)"
    v = a.answer_value
    if isinstance(v, list):
        return ", ".join(str(x) for x in v)
    return str(v)


class QAPDF(FPDF):
    def __init__(self, *, footer_tag: str) -> None:
        super().__init__()
        self._footer_tag = (footer_tag or "")[:100]
        self.set_auto_page_break(auto=True, margin=18)
        self.set_margins(18, 16, 18)
        self.set_creator("GrantFiller")

    def header(self) -> None:
        self.set_font("Helvetica", "B", 14)
        self.set_text_color(32, 32, 32)
        self.cell(0, 8, _txt("Application responses"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, _txt("Responses to application questions"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.ln(2)
        self.set_text_color(0, 0, 0)

    def footer(self) -> None:
        self.set_y(-16)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(120, 120, 120)
        tag = _txt(self._footer_tag) if self._footer_tag else ""
        if tag:
            self.cell(0, 5, f"Page {self.page_no()}  ·  {tag}", align="C")
        else:
            self.cell(0, 5, f"Page {self.page_no()}", align="C")


def build_qa_pdf(
    grant: Grant,
    questions: list[Question],
    answers: list[Answer],
    context: ExportContext | None = None,
) -> bytes:
    """Q&A PDF: no internal review / needs-input tags in body (product requirement)."""
    ctx = context or _default_context()
    answer_map = {a.question_id: a for a in answers}
    tag = (grant.name or "Grant")[:100]
    pdf = QAPDF(footer_tag=tag)
    pdf.set_title(_txt(f"{grant.name} — application responses"))
    pdf.add_page()
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(40, 40, 40)
    pdf.multi_cell(0, 6, _txt(f"Grant: {grant.name}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if ctx.organization_line:
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(70, 70, 70)
        pdf.multi_cell(0, 5, _txt(ctx.organization_line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(40, 40, 40)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(80, 80, 80)
    pdf.multi_cell(0, 5, _txt(ctx.exported_at_label), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)
    for i, q in enumerate(sorted(questions, key=lambda x: x.sort_order), start=1):
        y0 = pdf.get_y()
        if i > 1 and y0 < pdf.h - 35:
            pdf.set_draw_color(220, 225, 230)
            w = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.line(pdf.l_margin, y0, pdf.l_margin + w, y0)
            pdf.ln(1)
        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(0, 6, _txt(f"Question {i}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _txt(q.question_text or ""), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)
        a = answer_map.get(q.question_id)
        val = _format_answer_value(a)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _txt(f"Answer: {val}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)
    out = BytesIO()
    pdf.output(out)
    return out.getvalue()


def _docx_bootstrap_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)


def build_qa_docx(
    grant: Grant,
    questions: list[Question],
    answers: list[Answer],
    context: ExportContext | None = None,
) -> bytes:
    """Word-compatible .docx for Q&A (polished for Word / Google Docs)."""
    ctx = context or _default_context()
    answer_map = {a.question_id: a for a in answers}
    doc = Document()
    _docx_bootstrap_styles(doc)
    core = doc.core_properties
    core.title = (grant.name or "Grant")[:200]
    core.subject = "Grant application — responses to questions"
    core.keywords = "grant application"
    t = doc.add_paragraph()
    t_run = t.add_run((grant.name or "Grant").strip() or "Grant")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    t.alignment = 0

    if ctx.organization_line:
        s = doc.add_paragraph()
        s_run = s.add_run(ctx.organization_line)
        s_run.italic = True
        s_run.font.size = Pt(10)
        s_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    m_run = meta.add_run(ctx.exported_at_label)
    m_run.italic = True
    m_run.font.size = Pt(9)
    m_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for i, q in enumerate(sorted(questions, key=lambda x: x.sort_order), start=1):
        doc.add_heading(f"Question {i}", level=2)
        q_p = doc.add_paragraph()
        text = (q.question_text or "").strip()
        if text:
            for j, line in enumerate(text.splitlines()):
                if j:
                    q_p.add_run().add_break()
                q_p.add_run(line)
        else:
            q_p.add_run(" ")
        for r in q_p.runs:
            r.font.size = Pt(11)
        a = answer_map.get(q.question_id)
        ans_p = doc.add_paragraph()
        ax = ans_p.add_run("Answer: ")
        ax.bold = True
        ax.font.size = Pt(11)
        ans_p.add_run(_format_answer_value(a)).font.size = Pt(11)
        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_qa_markdown(
    grant: Grant,
    questions: list[Question],
    answers: list[Answer],
    context: ExportContext | None = None,
) -> str:
    ctx = context or _default_context()
    answer_map = {a.question_id: a for a in answers}
    lines = [f"# {grant.name}", ""]
    if ctx.organization_line:
        lines.append(f"_{ctx.organization_line}_")
        lines.append("")
    lines.append(f"_{ctx.exported_at_label}_")
    lines.append("")
    for i, q in enumerate(sorted(questions, key=lambda x: x.sort_order), start=1):
        lines.append(f"## Question {i}")
        lines.append("")
        lines.append(q.question_text or "")
        lines.append("")
        a = answer_map.get(q.question_id)
        val = _format_answer_value(a)
        lines.append(f"**Answer:** {val}")
        lines.append("")
    return "\n".join(lines)
